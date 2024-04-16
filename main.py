# import docx NOT python-docx 
import docx 
import pandas as pd
from docx.shared import Pt
from docx.shared import Pt
from docx.shared import Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===========================================
# data retrieved from the excel file

def valuef1(i, j):
    if all(pd.notnull(df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,59:60].values)):
        return df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,59:60].values
    return df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,60:61].values

def valuevis(i, j):
    if all(pd.notnull(df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,45:46].values)):
        return df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,45:46].values
    return df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,69:70].values


def get_value(row):
    row[59] = row[59].replace(' ', '')
    if row[59] != '':    
        print(row[59])
        return row[59]
    
    return row[60]

def get_value2(row):
    row[45] = row[45].replace(' ', '')
    if row[45] != '':
        
        return row[45]
    return row[69]
        
def add_page_number_footer(doc):
    section = doc.sections[0]  # Assuming there's only one section
    footer = section.footer
    footer.paragraphs[0].clear()  # Clear existing content if any
    
    # Create the paragraph for the page number
    page_number_paragraph = footer.paragraphs[0]
    page_number_run = page_number_paragraph.add_run()
    page_number_run.add_text("Page ")
    page_field = OxmlElement('w:fldSimple')
    page_field.set(qn('w:instr'), 'PAGE')
    page_number_run._r.append(page_field)
    page_number_run.add_text(" of ")
    total_pages_field = OxmlElement('w:fldSimple')
    total_pages_field.set(qn('w:instr'), 'NUMPAGES')
    page_number_run._r.append(total_pages_field)
    page_number_run.font.size = Pt(10)
    page_number_run.font.name = 'Times New Roman'
    page_number_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER



def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ["top", "start", "bottom", "end"]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)
    
def process(file): 
    
    global df
    df = pd.read_excel(file)
    df = df[df.iloc[:, 0].str.contains('\d', na=False)]

    doc = docx.Document() 
        
    suites = []
    for i in df.iloc[:, 0].unique():
        studies = []
        for j in df.iloc[:, 1].unique():
            studies.append(
                {
                    "suite": i,
                    "study": j,
                    "days": df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,3:4].values,
                    "condition": df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,4:5].values,
                    "appearance": df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,57:58].values,
                    "pH": df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,46:47].values,
                    "soluble_f1": df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].apply(get_value, axis=1).values,                
                    "TAMC": df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,81:82].values,
                    "flavour": df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,126:127].values,
                    "odour": df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,127:128].values,
                    "viscosity": df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].apply(get_value2, axis=1).values,
                    "dispensing": df[(df.iloc[:, 0] == i) & (df.iloc[:, 1] == j)].iloc[:,58:59].values
                }
            )   
        suites.append(studies)
            
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "\t\t Bussiness Use"
    header_para.runs[0].font.size = Pt(8)

    # add page 2 of 3 footer
    add_page_number_footer(doc)


    doc.styles['Normal'].font.name = 'Times New Roman'
    # set the font size of whole document
    doc.styles['Normal'].font.size = Pt(11)

    para1 = doc.add_paragraph()
    para1.add_run("Stability Data for ")
    para1.add_run("Crest Cavity Protection").font.highlight_color = 7
    para1.add_run(" Toothpaste Formulations\n")

    for run in para1.runs:
        run.bold = True
        
    para1.alignment = 1

    data = [
        ["Product Name", "Formula ID"],
        ["Crest Cavity Protection Herbal collection", "90466094"],
        ["Crest Cavity Protection Fresh Mint", "90473092"],
        ["Crest Cavity Protection Extra Fresh", "90456389"]
    ]

    table = doc.add_table(rows=4, cols=2)

    table.style = 'Table Grid'
    # list all the styles

    heading_row = table.rows[0]
    for cell in heading_row.cells:
        cell.background_color = 7
        cell.width = 100

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = data[i][j]
            cell.paragraphs[0].alignment = 1
            if i > 0:
                cell.paragraphs[0].runs[0].font.highlight_color = 7
            else:
                cell.paragraphs[0].runs[0].bold = True
                

    para2 = doc.add_paragraph()
    para2.add_run("\n\n\nSummary").bold = True
    para2.add_run("\n\nThe formulations ")
    para2.add_run("90466094, 90473092 and 90456389").font.highlight_color = 7
    para2.add_run(
    " belong to a family or ‘cluster’ of dentifrice products because the base chassis \
    is equivalent and made by the same manufacturing process and with the \
    same raw materials. The only differences between the cluster members \
    are in the flavours and visual attributes that make up a very small \
    percentage of the individual formulations and so will not impact stability.\n\n\
    The shelf-life was assigned by 6 months 40°C/75% RH accelerated and 36 \
    months 30°C/75% RH long-term stability on a representative formulation \
    from the cluster. Therefore, all the products from the cluster are \
    assigned a 36-month shelf-life based on the data from the representative \
    formulation. Data from the representative formulation is provided in \
    Results Tables on page 2."
    )

    para2.add_run("\n\n\nConclusion").bold = True
    para2.add_run(
    "\n\nFor the specific cluster member formulations, confirmatory stability \
    will be generated at the start of commercial production to include \
    formulas. The testing will cover sensory, chemical, physical and micro-biological \
    requirements for the shelf life of the product, when stored under 30°C/75% RH \
    long term stability conditions. ")
    para2.add_run("\n\n\n\n\nMengnan Ju, OC Analytical")

    para2.alignment = 3

    doc.add_page_break()



    for i in suites:
        for j in i:
            doc.add_paragraph('\n')
            rows = len(j['appearance'])
            
            condition = j['condition']
            t5, t6 = 0, 0
            for x in condition:
                if '40' in x[0]:
                    t6 += 1
                else:  
                    t5 += 1
                
            if rows < 1:
                continue
            
            table = doc.add_table(rows=8+rows, cols=10)
            table.cell(0, 0).merge(table.cell(0, 8))
            for i in range(6):
                table.cell(i, 0).merge(table.cell(i , 1))
            for i in range(8):
                table.cell(5, i+2).merge(table.cell(6 , i+2))

            table.cell(8, 0).merge(table.cell(7+t5, 0))
            table.cell(8+t5, 0).merge(table.cell(7+rows, 0))

            # Set column widths as per the image structure
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    if (i in [1, 2, 3, 5]):
                        set_cell_margins(cell, top=0, start=50, bottom=0, end=50)  # adjust padding as needed
                    # Set font size as needed
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            run.font.size = 6  # Set font size as needed
            table.style = 'Table Grid'  # Use an appropriate style

            table.cell(0, 0).text = f'Suite {j["suite"]}, Study {j["study"]}, Representative Cluster Member (97221263)'
            table.cell(0, 0).paragraphs[0].runs[0].bold = True
            table.cell(0, 0).paragraphs[0].runs[0].alignment = 1
            table.cell(0, 0).paragraphs[0].runs[0].font.highlight_color = 7

            table.cell(1, 1).text = 'Parameter '
            table.cell(1, 1).paragraphs[0].runs[0].bold = True
            table.cell(1, 2).text = 'Appearance'
            table.cell(1, 3).text = 'pH'
            table.cell(1, 4).text = 'Soluble Fluoride'
            table.cell(1, 5).text = 'Total Aerobic Microbial Count'
            table.cell(1, 6).text = 'Flavour'
            table.cell(1, 7).text = 'Odour'
            table.cell(1, 8).text = 'Viscosity'
            table.cell(1, 9).text = 'Dispensing'

            table.cell(2, 1).text = 'e.g. Parameter list ID in GAIMS'
            table.cell(2, 1).paragraphs[0].runs[0].bold = True
            table.cell(2, 2).text = 'NA-HCB-L30'
            table.cell(2, 3).text = 'NA-HCB-E08'
            table.cell(2, 4).text = 'NA-HCB-M32 or NA-HCB-M35'
            table.cell(2, 5).text = 'NB-HCX-N03A'
            table.cell(2, 6).text = 'NC-HCX-D23S'
            table.cell(2, 7).text = 'NC-HCX-D23S'
            table.cell(2, 8).text = 'NA-HCB-D61 or  NA-HCB-D60'
            table.cell(2, 9).text = 'NA-HCB-L32'

            table.cell(3, 1).text = 'Parameter list ID in Nexus'
            table.cell(3, 1).paragraphs[0].runs[0].bold = True
                
            table.cell(4, 1).text = 'Units'
            table.cell(4, 1).paragraphs[0].runs[0].bold = True
            table.cell(4, 2).text = 'Pass/Fail'
            table.cell(4, 3).text = 'N/A'
            table.cell(4, 4).text = 'ug/g'
            table.cell(4, 5).text = 'Pass/Fail'
            table.cell(4, 6).text = 'Pass/Fail'
            table.cell(4, 7).text = 'Pass/Fail'
            table.cell(4, 8).text = 'bku'

            table.cell(5, 1).text = 'Stabiltiy Specification'
            table.cell(5, 1).paragraphs[0].runs[0].bold = True
            table.cell(5, 2).text = '*In A white to off-white paste that meets all texture requirements'
            table.cell(5, 3).text = '5.5-10.0'
            table.cell(5, 4).text = '≥ 500'
            table.cell(5, 5).text = 'Total Aerobic Microbes ≤100'
            table.cell(5, 6).text = 'Pass/Fail'
            table.cell(5, 7).text = 'Pass/Fail'
            table.cell(5, 8).text = '≤ 180'

            table.cell(6, 0).text = 'Condition'
            table.cell(6, 0).paragraphs[0].runs[0].bold = True
            table.cell(6, 1).text = 'Time Point (Days)'
            table.cell(6, 1).paragraphs[0].runs[0].bold = True

            table.cell(7, 0).text = 'Initial'
            table.cell(8, 0).text = '30C/75%RH'
            table.cell(8+t5, 0).text = '40C/75%RH'

            
            for x in range(rows):
                if 'initial' in condition[x][0].lower():
                    table.cell(7, 1).text = str(j['days'][x][0])
                    table.cell(7, 2).text = str(j['appearance'][x][0])
                    table.cell(7, 3).text = str(j['pH'][x][0])
                    table.cell(7, 4).text = str(j['soluble_f1'][x])
                    table.cell(7, 5).text = str(j['TAMC'][x][0])
                    table.cell(7, 6).text = str(j['flavour'][x][0])
                    table.cell(7, 7).text = str(j['odour'][x][0])
                    table.cell(7, 8).text = str(j['viscosity'][x])
                    table.cell(7, 9).text = str(j['dispensing'][x][0])
            
            i = 1
            for x in range(rows):
                if not '40' in condition[x][0] and not 'initial' in condition[x][0].lower():
                    table.cell(i+7, 1).text = str(j['days'][x][0])
                    table.cell(i+7, 2).text = str(j['appearance'][x][0])
                    table.cell(i+7, 3).text = str(j['pH'][x][0])
                    table.cell(i+7, 4).text = str(j['soluble_f1'][x])
                    table.cell(i+7, 5).text = str(j['TAMC'][x][0])
                    table.cell(i+7, 6).text = str(j['flavour'][x][0])
                    table.cell(i+7, 7).text = str(j['odour'][x][0])
                    table.cell(i+7, 8).text = str(j['viscosity'][x])
                    table.cell(i+7, 9).text = str(j['dispensing'][x][0])
                    i+=1
            i = 1
            for x in range(rows):
                if '40' in condition[x][0]:
                    table.cell(i+7+t5, 1).text = str(j['days'][x][0])
                    table.cell(i+7+t5, 2).text = str(j['appearance'][x][0])
                    table.cell(i+7+t5, 3).text = str(j['pH'][x][0])
                    table.cell(i+7+t5, 4).text = str(j['soluble_f1'][x])
                    table.cell(i+7+t5, 5).text = str(j['TAMC'][x][0])
                    table.cell(i+7+t5, 6).text = str(j['flavour'][x][0])
                    table.cell(i+7+t5, 7).text = str(j['odour'][x][0])
                    table.cell(i+7+t5, 8).text = str(j['viscosity'][x])
                    table.cell(i+7+t5, 9).text = str(j['dispensing'][x][0])
                    i += 1
                # font color to green
                # table.cell(i+7, j+1).paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(0, 128, 0)
                

    doc.add_paragraph('\n')  # Add a paragraph after the table if needed


    # change font size of all tables in document
    for table in doc.tables[1:]:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = 1
                        


    para3 = doc.add_paragraph()
    para3.add_run("*The description of the visual appearance may differ between cluster members\n \
    **Data prior to 90 days (3 months), has not been reported due to an experimental issue. \
    Subsequent data is all within specification indicating good stability.\n\
    NS – not scheduled ")
    para3.alignment = 1

    doc.save('document.docx') 



import streamlit as st
import pandas as pd


def main():
    st.title("Upload xlsx File and Generate")

    # Upload Xlsx file
    uploaded_file = st.file_uploader("Upload Xlsx file", type=['xlsx'])

    if uploaded_file is not None:
        st.success("File uploaded successfully.")
        
        st.write(uploaded_file.name)

        # Generate button to process the uploaded file
        if st.button("Generate"):
            # wrote processing...
            st.write("Processing...")
            process(uploaded_file)
            # write processing complete and delete previous line
            st.write("Processing complete.")
            
            
            

if __name__ == "__main__":
    main()
